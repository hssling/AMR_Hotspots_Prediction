import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_latex(text):
    # For safety, remove $ delimiters if they exist around pure numbers/symbols
    # But mainly rely on unicode replacement
    text = text.replace('^2', '²').replace('^3', '³')
    text = text.replace(r'\times', '×')
    text = text.replace(r'\le', '≤').replace(r'\ge', '≥')
    text = text.replace('$', '')
    return text

def add_formatted_text(paragraph, text):
    # Pre-process math
    if '$' in text:
        text = clean_latex(text)
        
    # Regex to capture:
    # 1. Bold: **text**
    # 2. Italic: *text*
    # 3. Superscript: ^text^ or ^1 (naive)
    
    # 1. Split by Bold
    tokens = re.split(r'(\*\*.*?\*\*)', text)
    for token in tokens:
        is_bold = False
        content = token
        if token.startswith('**') and token.endswith('**'):
            is_bold = True
            content = token[2:-2]
            
        # 2. Split by Italic
        sub_tokens = re.split(r'(\*.*?\*)', content)
        for sub_token in sub_tokens:
            is_italic = False
            sub_content = sub_token
            if sub_token.startswith('*') and sub_token.endswith('*'):
                is_italic = True
                sub_content = sub_token[1:-1]
                
            # 3. Handle Superscripts defined as $^N$, ^N, or [^N]
            # Pattern: matches $^12$ or ^12 or [^12]
            # We treat [^N] as a pure citation -> superscript N (no brackets)
            sup_pattern = r'(\$\^[\d,–-]+\$|\[\^[\d,–-]+\]|\^[\d,–-]+)'
            
            sup_tokens = re.split(sup_pattern, sub_content)
            for atom in sup_tokens:
                if not atom: continue
                
                is_sup = False
                atom_text = atom
                
                # Check match
                if (atom.startswith('$^') and atom.endswith('$')) or atom.startswith('^'):
                    is_sup = True
                    atom_text = atom.replace('$','').replace('^','')
                elif atom.startswith('[^') and atom.endswith(']'):
                    is_sup = True
                    atom_text = atom[2:-1] # Remove [^ and ]
                    
                run = paragraph.add_run(atom_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = is_bold
                run.italic = is_italic
                if is_sup:
                    run.font.superscript = True

def create_manuscript_package(source_md, output_dir, figure_base):
    print("Generating High-Precision Manuscript (Audit Fixes with Image Patch)...")
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    
    with open(source_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    tables_buffer = []
    current_table = []
    in_table = False
    table_title_buffer = None
    
    for line in lines:
        line = line.strip()
        
        # Table Logic
        if line.startswith('**Table'):
            table_title_buffer = line
            continue
        if line.startswith('|'):
            in_table = True
            current_table.append(line)
            continue
        if in_table and not line.startswith('|'):
            if current_table:
                tables_buffer.append({'title': table_title_buffer, 'rows': current_table})
            current_table = []
            in_table = False
            table_title_buffer = None
            if not line: continue
            
        if not line:
            if len(doc.paragraphs) > 0: doc.add_paragraph()
            continue
            
        # Headers
        if line.startswith('## '):
            doc.add_page_break()
            p = doc.add_heading(line.strip('# '), 2)
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.color.rgb = RGBColor(0,0,0)
        elif line.startswith('# '):
            p = doc.add_heading(line.strip('# '), 1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.color.rgb = RGBColor(0,0,0)
            
        # Images
        elif line.startswith('![') and '](' in line:
            try:
                start_path = line.find('](') + 2
                end_path = line.find(')', start_path)
                img_path = line[start_path:end_path]
                
                # Robust Search Paths
                search_paths = [
                    figure_base,
                    r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction\outputs\figures",
                    r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction\outputs\advanced_analytics",
                    r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction\outputs\scorecard",
                    r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction\outputs\manuscript_figures"
                ]
                
                final_img_path = None
                # Try relative resolution first
                if not os.path.isabs(img_path):
                     # Try combining with base first
                     candidate = os.path.join(figure_base, img_path)
                     if os.path.exists(candidate): final_img_path = candidate
                else: 
                     if os.path.exists(img_path): final_img_path = img_path
                     
                if not final_img_path:
                    # Fallback search
                    for search_dir in search_paths:
                        candidate = os.path.join(search_dir, os.path.basename(img_path))
                        if os.path.exists(candidate):
                            final_img_path = candidate
                            break
                
                if final_img_path:
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(final_img_path, width=Inches(6))
                        print(f"SUCCESS: Embedded Image: {final_img_path}")
                    except Exception as e:
                        print(f"CRITICAL ERROR adding picture {final_img_path}: {e}")
                        raise e
                else:
                    print(f"FAILURE: Image not found: {img_path}")
                    print(f"  Searched in: {figure_base} and fallback paths.")
                    # raise Exception(f"Missing Image: {img_path}") # Uncomment to force crash
                    p = doc.add_paragraph(f"[IMAGE MISSING: {img_path}]")
                    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            except Exception as e:
                print(f"Error adding image: {e}")
            
        # Text
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            
    # Flush last table
    if in_table and current_table:
        tables_buffer.append({'title': table_title_buffer, 'rows': current_table})
        
    # Append Tables at the END (After References)
    if tables_buffer:
        doc.add_page_break()
        h = doc.add_heading("Tables", 1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.runs[0].font.name = 'Times New Roman'
        h.runs[0].font.color.rgb = RGBColor(0,0,0)
        
        for tbl in tables_buffer:
            doc.add_paragraph()
            if tbl['title']:
                p = doc.add_paragraph()
                add_formatted_text(p, tbl['title'])
            
            rows = [r for r in tbl['rows'] if '---' not in r]
            if rows:
                cols = len(rows[0].split('|')) - 2
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for i, row_str in enumerate(rows):
                    cells = row_str.split('|')[1:-1]
                    for j, txt in enumerate(cells):
                        if j < len(table.columns):
                            cell_p = table.cell(i, j).paragraphs[0]
                            add_formatted_text(cell_p, txt.strip())
                            if i == 0:
                                for run in cell_p.runs: run.font.bold = True
                                
    out_file = os.path.join(output_dir, "Manuscript_1_IJMR_Comprehensive_V2.docx")
    doc.save(out_file)
    print(f"Saved: {out_file}")

if __name__ == "__main__":
    base_dir = r"C:\Users\hssli\.gemini\antigravity\brain\90c42530-5be5-49cb-a7b5-e960c5582f78"
    out_dir = r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction\submission"
    create_manuscript_package(
        os.path.join(base_dir, "Manuscript_AMR_Hotspots_IJMR.md"),
        out_dir,
        base_dir
    )
