
import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_simple(source_md, output_filename, figure_base):
    print(f"Generating {output_filename}...")
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    with open(source_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_lines = []

    for line in lines:
        line = line.strip()
        if not line: continue

        if line.startswith('# '):
            doc.add_heading(line.strip('# '), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.strip('# '), level=2)
        elif line.startswith('![') and 'figure' in line.lower():
            # Image Handler
            # Expecting ![Alt](figures_manuscript2/fig2.png)
            match = re.search(r'\((.*?)\)', line)
            if match:
                rel_path = match.group(1)
                # Cleanup path syntax if valid file URI
                if rel_path.startswith('file:///'):
                     rel_path = rel_path.replace('file:///', '')
                
                # If path isn't absolute, join
                if not os.path.isabs(rel_path):
                    full_path = os.path.abspath(rel_path)
                else:
                    full_path = rel_path
                    
                # Fix specific path issue for this run
                if 'figures_manuscript2' in full_path and not os.path.exists(full_path):
                     full_path = os.path.join(r'd:\research-automation\TB multiomics\AMR_Hotspots_Prediction', rel_path)

                if os.path.exists(full_path):
                    try:
                        doc.add_picture(full_path, width=Inches(5.0))
                    except:
                        doc.add_paragraph("[Image Placeholder]")
                        
        # 3. Tables (Simple Markdown Table Parser)
        elif line.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
        
        # 4. Standard Text
        else:
            if in_table:
                # Process the accumulated table
                try:
                    valid_rows = [row for row in table_lines if '---' not in row]
                    if valid_rows:
                        cols = len(valid_rows[0].split('|')) - 2
                        current_table = doc.add_table(rows=len(valid_rows), cols=cols)
                        current_table.style = 'Table Grid'
                        
                        for i, row_md in enumerate(valid_rows):
                            cells = row_md.split('|')[1:-1]
                            for j, cell_text in enumerate(cells):
                                current_table.cell(i, j).text = cell_text.strip()
                except Exception as e:
                    print(f"Table parsing error: {e}")
                
                in_table = False
                table_lines = []
                
            doc.add_paragraph(line)

    doc.save(output_filename)

if __name__ == "__main__":
    base = r"C:\Users\hssli\.gemini\antigravity\brain\90c42530-5be5-49cb-a7b5-e960c5582f78"
    # 1. Main Manuscript
    create_docx_simple(
        os.path.join(base, "Manuscript_2_Decoupling_AMR.md"),
        "submission_manuscript2/Manuscript_2_IJCCM_Final_Submission.docx",
        "outputs/figures_manuscript2"
    )
    
    # 2. Cover Letter
    create_docx_simple(
        os.path.join(base, "Cover_Letter_IJCCM.md"),
        "submission_manuscript2/Cover_Letter_IJCCM.docx",
        ""
    )
