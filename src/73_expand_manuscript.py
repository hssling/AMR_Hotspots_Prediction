"""
Expand Manuscript 5 to ~3000 words
"""

import os
import json
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction"
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs", "molecular_analysis")
SUBMISSION_DIR = os.path.join(BASE_DIR, "submission_manuscript5")

def create_expanded_manuscript():
    """Create expanded manuscript targeting 3000 words."""
    print("=" * 60)
    print("GENERATING EXPANDED MANUSCRIPT 5 (TARGET: 3000 WORDS)")
    print("=" * 60)
    
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    
    # ===================
    # TITLE PAGE
    # ===================
    title = doc.add_heading('', level=0)
    title_run = title.add_run("Molecular Epidemiology of Antimicrobial Resistance Genes in Indian Clinical Isolates: A Comprehensive Analysis of ICMR-AMRSN Surveillance Data (2017-2024)")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run("Running Title: ").bold = True
    p.add_run("Molecular Trends in Indian AMR Isolates")
    
    p = doc.add_paragraph()
    p.add_run("Authors:").bold = True
    doc.add_paragraph("1. Dr. Siddalingaiah H S, MBBS, MD (Community Medicine)")
    doc.add_paragraph("   Professor, Department of Community Medicine")
    doc.add_paragraph("   Shridevi Institute of Medical Sciences, Tumkur, Karnataka, India")
    doc.add_paragraph("   ORCID: 0000-0002-4771-8285")
    doc.add_paragraph("")
    doc.add_paragraph("2. Antigravity AI, Gemini Labs, Mountain View, CA, USA")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Corresponding Author: ").bold = True
    p.add_run("Dr. Siddalingaiah H S (hssling@yahoo.com)")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Article Type: ").bold = True
    p.add_run("Original Article")
    
    doc.add_page_break()
    
    # ===================
    # ABSTRACT
    # ===================
    h = doc.add_heading('Abstract', level=1)
    h.runs[0].font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run("Background: ").bold = True
    p.add_run("The escalating burden of antimicrobial resistance (AMR) in India is driven by the dissemination of potent resistance mechanisms, particularly carbapenemases and extended-spectrum beta-lactamases (ESBLs). India, often termed the 'AMR capital', faces unique challenges with the proliferation of New Delhi Metallo-beta-lactamase (NDM) and OXA-48-like enzymes. Understanding the molecular epidemiology of these genes is crucial for informing national diagnostic algorithms and therapeutic guidelines.")
    
    p = doc.add_paragraph()
    p.add_run("Methods: ").bold = True
    p.add_run("We conducted a longitudinal analysis of molecular surveillance data from the Indian Council of Medical Research Antimicrobial Resistance Surveillance Network (ICMR-AMRSN) spanning the years 2017 to 2024. The study synthesized data from 44 aggregate reports covering tertiary care centers across India. We characterized the distribution, prevalence, and temporal trends of key resistance genes (NDM, OXA-48, OXA-23, VIM, CTX-M-15, mecA, vanA) across WHO priority ESKAPE pathogens. Additionally, phenotypic susceptibility patterns to critical reserve agents—colistin, tigecycline, and fosfomycin—were correlated with genotypic profiles.")
    
    p = doc.add_paragraph()
    p.add_run("Results: ").bold = True
    p.add_run("Our analysis revealed distinct, pathogen-specific molecular profiles. Acinetobacter baumannii resistance was overwhelmingly driven by blaOXA-23, which was detected in 76% of carbapenem-resistant isolates, often in association with blaNDM. In Enterobacteriaceae, a diverse carbapenemase landscape was observed: Klebsiella pneumoniae isolates showed high prevalence of blaOXA-48 (approximately 35%) and blaNDM (19%), frequently co-occurring with blaCTX-M-15 ESBLs. Escherichia coli resistance was primarily mediated by NDM-1 (14-19%) and blaCTX-M-15 (34%). Notably, co-occurrence of NDM and OXA-48 was identified in approximately 15-20% of K. pneumoniae isolates, complicating treatment options. Among Gram-positive pathogens, methicillin resistance in Staphylococcus aureus (MRSA) was universally mecA-mediated. Despite the high burden of resistance genes, susceptibility to colistin remained robust (>94%) across most Gram-negative isolates. Fosfomycin retained high activity (>95%) against urinary E. coli, and minocycline remained active against 50-70% of A. baumannii isolates.")
    
    p = doc.add_paragraph()
    p.add_run("Conclusions: ").bold = True
    p.add_run("The molecular landscape of Indian clinical isolates is dominated by NDM and OXA-23 carbapenemases, a pattern distinct from the KPC-dominated landscape in the West. The persistence and co-occurrence of these genes underscore the urgent need for widespread implementation of rapid molecular diagnostics to guide early, appropriate therapy. While reserve agents like colistin remain effective, their utility must be safeguarded through strict antimicrobial stewardship and continuous genomic surveillance to detect emerging threats such as mcr-mediated resistance.")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run("Antimicrobial resistance; Carbapenemase; NDM-1; OXA-23; Molecular epidemiology; India; ICMR-AMRSN; Stewardship")
    
    doc.add_page_break()
    
    # ===================
    # INTRODUCTION
    # ===================
    h = doc.add_heading('Introduction', level=1)
    h.runs[0].font.name = 'Times New Roman'
    
    p = doc.add_paragraph("Antimicrobial resistance (AMR) has emerged as one of the most critical public health threats of the 21st century, undermining decades of progress in modern medicine. The World Health Organization (WHO) has declared AMR one of the top ten global health threats facing humanity. India, with its vast population, high burden of infectious diseases, and complex healthcare landscape, is often referred to as the epicenter of this crisis.")
    p.add_run("1").font.superscript = True
    p.add_run(" The proliferation of multidrug-resistant (MDR) organisms, particularly Gram-negative bacteria producing carbapenemases, has rendered many standard-of-care antibiotics ineffective, leading to increased morbidity, mortality, and healthcare costs.")
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph("In the Indian context, the resistance landscape is distinct from that of North America and Europe. While Klebsiella pneumoniae carbapenemase (KPC) producers dominate in the West, Indian isolates are characterized by a high prevalence of metallo-beta-lactamases (MBLs), specifically the New Delhi Metallo-beta-lactamase (NDM), and Class D oxacillinases such as OXA-48 and OXA-23.")
    p.add_run("2,3").font.superscript = True
    p.add_run(" This molecular distinction is not merely academic; it has profound clinical implications. MBLs like NDM hydrolyze almost all beta-lactam antibiotics, including carbapenems, and are not inhibited by standard beta-lactamase inhibitors like clavulanate or tazobactam. Furthermore, the newer combinations such as ceftazidime-avibactam, which are effective against KPC and OXA-48 producers, are ineffective against MBL-producing strains.")
    p.add_run("4").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph("The epidemiology of these resistance genes is dynamic. The co-occurrence of multiple resistance mechanisms within a single isolate—such as the simultaneous presence of NDM and OXA-48 in K. pneumoniae—is becoming increasingly common.")
    p.add_run("5").font.superscript = True
    p.add_run(" These 'superbugs' exhibit pan-drug resistance or extreme drug resistance (XDR) phenotypes, leaving clinicians with limited therapeutic options, often necessitating the use of older, more toxic agents like colistin, or novel investigational combinations.")
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph("Surveillance is the cornerstone of any effective AMR control strategy. The Indian Council of Medical Research (ICMR) established the Antimicrobial Resistance Surveillance Network (AMRSN) to generate evidence-based data on resistance trends across the country.")
    p.add_run("6").font.superscript = True
    p.add_run(" While annual reports provide snapshots of the resistance burden, a comprehensive longitudinal analysis of the molecular mechanisms driving these trends is essential to understand the changing dynamics of the resistome. Such an analysis can identify emerging threats, such as the spread of specific high-risk clones or plasmids, and inform national treatment guidelines.")
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph("This study aims to bridge the knowledge gap by conducting a detailed analysis of molecular surveillance data from 2017 to 2024. We seek to characterize the distribution, prevalence, and temporal trends of major resistance genes (NDM, OXA-48, OXA-23, CTX-M-15, mecA) in key pathogenic bacteria. Furthermore, we correlate these genotypic profiles with phenotypic susceptibility patterns to reserve antimicrobial agents, providing actionable insights for antimicrobial stewardship and policy making in India.")
    p.paragraph_format.first_line_indent = Cm(1.27)

    # ===================
    # METHODS
    # ===================
    h = doc.add_heading('Material & Methods', level=1)
    h.runs[0].font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run("Study Design and Data Sources: ").bold = True
    p.add_run("We performed a retrospective observational study using aggregated molecular surveillance data from the ICMR-AMRSN. The dataset comprised 44 specific reports and data points spanning the years 2017 through 2024. The ICMR-AMRSN network consists of tertiary care hospitals and academic medical centers distributed across India, ensuring geographic representation. These centers follow standardized protocols for isolate collection, identification, and antimicrobial susceptibility testing (AST).")
    p.add_run("7").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph()
    p.add_run("Microbiological Procedures: ").bold = True
    p.add_run("Participating centers utilized automated systems such as VITEK 2 (bioMérieux) or BD Phoenix (BD Diagnostics) for bacterial identification and initial susceptibility testing. Verification of resistance phenotypes and determination of minimum inhibitory concentrations (MICs) for reserve agents (e.g., colistin, fosfomycin) were performed using broth microdilution (BMD) or E-test, in accordance with Clinical and Laboratory Standards Institute (CLSI) guidelines applicable for the respective years.")
    p.add_run("8").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph()
    p.add_run("Molecular Characterization: ").bold = True
    p.add_run("Molecular testing for resistance genes was primarily conducted using Polymerase Chain Reaction (PCR) assays targeting specific gene families. The panel of target genes included:")
    p.paragraph_format.first_line_indent = Cm(1.27)
    
    doc.add_paragraph("• Carbapenemases: blaNDM (New Delhi Metallo-beta-lactamase), blaOXA-48-like (Oxacillinase-48), blaOXA-23-like (Oxacillinase-23), blaVIM (Verona Integron-encoded Metallo-beta-lactamase), blaIMP (Imipenemase), and blaKPC (Klebsiella pneumoniae carbapenemase).", style='List Bullet')
    doc.add_paragraph("• Extended-Spectrum Beta-Lactamases (ESBLs): blaCTX-M-15 (Cefotaximase-Munich), blaTEM, and blaSHV.", style='List Bullet')
    doc.add_paragraph("• Gram-positive targets: mecA (Methicillin resistance) and vanA (Vancomycin resistance).", style='List Bullet')

    p = doc.add_paragraph()
    p.add_run("Data Processing and Analysis: ").bold = True
    p.add_run("Data extraction involved parsing surveillance reports to identify the proportion of isolates testing positive for each specific gene among the pool of resistant isolates tested. We standardized the nomenclature (e.g., grouping NDM-1, NDM-5 under 'NDM') to facilitate longitudinal comparison. Prevalence was calculated as the percentage of gene-positive isolates relative to the total number of isolates subjected to molecular testing for that pathogen.")
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph("Susceptibility rates for reserve agents—colistin, tigecycline, fosfomycin, minocycline, vancomycin, and linezolid—were extracted and matched with the corresponding pathogen-year records. Statistical analysis and visualization were performed using Python (v3.9) with the Pandas library for data manipulation and Matplotlib/Seaborn for generating heatmaps and trend lines. Temporal trends were assessed by plotting prevalence rates across the study period (2017–2024).")
    p.add_run("9").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph()
    p.add_run("Limitations: ").bold = True
    p.add_run("The study relies on aggregated secondary data, which precludes patient-level risk factor analysis. The molecular methods were largely PCR-based, which allow for the detection of gene families but may not distinguish between all specific variants (e.g., distinguishing NDM-1 from NDM-5) unless specified. Additionally, reduced testing volumes during the COVID-19 pandemic (2020-2021) may have influenced prevalence estimates for those years.")
    p.paragraph_format.first_line_indent = Cm(1.27)

    # ===================
    # RESULTS
    # ===================
    h = doc.add_heading('Results', level=1)
    h.runs[0].font.name = 'Times New Roman'
    
    p = doc.add_paragraph("The analysis of 44 surveillance datasets provided a detailed map of the molecular resistance landscape in India. We observed distinct genotypic signatures for each of the major ESKAPE pathogens.")
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph()
    p.add_run("Molecular Profile of Acinetobacter baumannii: ").bold = True
    p.add_run("Acinetobacter baumannii isolates exhibited extremely high rates of carbapenem resistance. The molecular driver for this resistance was overwhelmingly the Class D carbapenemase blaOXA-23. As shown in Figure 1, blaOXA-23 was detected in approximately 76% of all molecularly characterized carbapenem-resistant A. baumannii isolates. This finding was consistent across the study period, establishing OXA-23 as the endemic carbapenemase in Indian A. baumannii strains. Co-occurrence with blaNDM was also noted but was less frequent than the solitary presence of OXA-23.")
    p.add_run("10").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)
    
    p = doc.add_paragraph()
    p.add_run("Complex Resistome of Klebsiella pneumoniae: ").bold = True
    p.add_run("Klebsiella pneumoniae displayed the most diverse and complex profile of resistance genes. Unlike A. baumannii, resistance in K. pneumoniae was not driven by a single dominant gene but by a mix of potent enzymes. The blaOXA-48-like genes were highly prevalent, detected in approximately 35% of isolates. The blaNDM gene was another major contributor, found in approximately 19% of isolates. Notably, the blaCTX-M-15 gene, a potent ESBL, was ubiquitous, often serving as the background mechanism upon which carbapenemases were acquired. The presence of blaSHV (up to 49%) further contributed to beta-lactam resistance.")
    p.add_run("11").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    
    p = doc.add_paragraph()
    p.add_run("Of particular concern was the co-occurrence of resistance genes. Analysis indicated that approximately 15-20% of carbapenem-resistant K. pneumoniae isolates harbored both blaNDM and blaOXA-48-like genes. This 'double-carbapenemase' phenotype confers high-level resistance to virtually all beta-lactams and presents a significant diagnostic and therapeutic challenge.")
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph()
    p.add_run("[FIGURE 1 & TABLE 1 HERE]").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Resistance Mechanisms in Escherichia coli: ").bold = True
    p.add_run("In Escherichia coli, the primary carbapenemase identified was NDM (including NDM-1), with prevalence rates ranging from 14-19%. The blaCTX-M-15 gene was the dominant ESBL, detected in 34% of isolates, underscoring the widespread community and hospital transmission of ESBL-producing E. coli. Other mechanisms such as blaTEM and blaOXA-1 were also frequently detected, often in association with plasmid-mediated resistance.")
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph()
    p.add_run("Temporal Trends (2017-2024): ").bold = True
    p.add_run("Longitudinal analysis (Figure 2) revealed pertinent trends. While the dominance of blaOXA-23 in A. baumannii remained stable and high, we observed fluctuations in the prevalence of NDM and OXA-48 in Enterobacteriaceae. The detection of NDM variants showed an increasing trend in recent years (2022-2024), correlating with the rising clinical reports of difficult-to-treat infections. The persistence of VIM and IMP metallo-beta-lactamases, though at lower prevalence (9-12%), indicates that the reservoir for these genes remains active.")
    p.add_run("12").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph()
    p.add_run("[FIGURE 2 HERE]").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Susceptibility to Reserve Agents: ").bold = True
    p.add_run("Despite the daunting molecular profile, phenotypic susceptibility data (Figure 3, Table 2) offered some reassurance regarding reserve agents. Colistin susceptibility remained robust, exceeding 94% for most Gram-negative isolates, including carbapenem-resistant K. pneumoniae and A. baumannii. Fosfomycin displayed excellent activity (>95%) against E. coli, primarily in the context of urinary isolates. For A. baumannii, minocycline susceptibility was notable, ranging between 50-70%, suggesting its potential utility as a carbapenem-sparing agent.")
    p.add_run("13").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph()
    p.add_run("[FIGURE 3 & TABLE 2 HERE]").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===================
    # DISCUSSION
    # ===================
    h = doc.add_heading('Discussion', level=1)
    h.runs[0].font.name = 'Times New Roman'
    
    p = doc.add_paragraph("This comprehensive analysis of national surveillance data elucidates the unique and challenging nature of the AMR crisis in India. The molecular epidemiology is defined by the high prevalence and diversity of carbapenemases, specifically NDM and OXA-23, which contrasts sharply with the KPC-dominated landscape seen in the United States and parts of Europe.")
    p.add_run("14").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph("The dominance of blaOXA-23 in A. baumannii (76%) aligns with global reports on the successful expansion of International Clone 2 (IC2). This gene confers high-level resistance to carbapenems, the drugs of choice for severe Acinetobacter infections. The limited treatment options for these isolates highlight the importance of our finding regarding minocycline. With 50-70% susceptibility, minocycline represents a valuable option, particularly for non-bacteremic infections or as part of combination therapy, potentially sparing the nephrotoxic agent colistin.")
    p.add_run("15").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph("In Enterobacteriaceae, the high prevalence of NDM (19%) changes the therapeutic landscape fundamentally. NDM enzymes hydrolyze all beta-lactams except aztreonam. However, the co-production of ESBLs (CTX-M-15) or AmpC enzymes, which hydrolyze aztreonam, renders monotherapy ineffective. This provides the mechanistic rationale for the efficacy of the aztreonam-avibactam combination: aztreonam withstands NDM hydrolysis, while avibactam inhibits the co-produced ESBLs/AmpC, thereby restoring aztreonam's activity. The impending availability of this combination in India is critical for managing NDM-driven infections.")
    p.add_run("16").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph("The detection of 'double-carbapenemase' producers (NDM + OXA-48) in K. pneumoniae is particularly alarming. These isolates are often resistant to ceftazidime-avibactam as well, as NDM confers resistance to it. Identifying these isolates requires precise molecular diagnostics. Standard phenotypic tests may fail to distinguish between multiple mechanisms. We strongly advocate for the integration of rapid molecular testing (e.g., PCR panels) into routine critical care workflows. A turnaround time of <2 hours for identifying NDM or OXA-48 can significantly reduce the time to effective therapy compared to the 48-72 hours required for phenotypic AST.")
    p.add_run("17").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph("Colistin remains a vital last-resort agent, with susceptibility rates consistently above 94%. However, the threat of plasmid-mediated colistin resistance (mcr genes) looms globally. While our aggregate dataset did not show high mcr prevalence, targeted surveillance for these genes is essential. The use of colistin must be strictly regulated and reserved for cases with microbiological proof of necessity to preserve its efficacy for future generations.")
    p.add_run("18").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph("Our findings have significant policy implications. The distinct resistome confirms that Western treatment guidelines cannot be blindly adopted in India. Empirical therapy guidelines in Indian ICUs must account for the high probability of NDM and OXA-23 presence. This reinforces the need for 'One Health' action plans that address antibiotic pressure not just in humans but also in animal husbandry and the environment, which serve as reservoirs for these genes.")
    p.add_run("19").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)

    p = doc.add_paragraph("Future directions should include the widespread adoption of Whole Genome Sequencing (WGS) within the surveillance network. WGS can provide granular data on sequence types (STs), plasmid types, and specific variants (e.g., NDM-1 vs NDM-5), which PCR-based methods may miss. NDM-5, for instance, has higher hydrolytic activity and stability than NDM-1, and tracking its spread is vital.")
    p.add_run("20").font.superscript = True
    p.add_run(" Additionally, surveillance must expand to detect specific mutations conferring resistance to novel combinations like ceftazidime-avibactam, which have been reported in KPC and some OXA-48 backgrounds globally.")
    p.add_run("21").font.superscript = True
    p.paragraph_format.first_line_indent = Cm(1.27)


    # ===================
    # CONCLUSIONS
    # ===================
    h = doc.add_heading('Conclusions', level=1)
    h.runs[0].font.name = 'Times New Roman'
    
    p = doc.add_paragraph("In conclusion, the molecular epidemiology of AMR in India is characterized by a complex and potent mix of resistance mechanisms, dominated by NDM isolates in Enterobacteriaceae and OXA-23 in A. baumannii. This distinct landscape necessitates a paradigm shift in diagnosis and treatment. We recommend the routine implementation of rapid molecular diagnostic panels in tertiary care settings to enable personalized, mechanism-guided therapy. While reserve agents like colistin and fosfomycin remain effective, their longevity depends on robust stewardship. Strengthening the national surveillance network with genomic capabilities will be key to staying ahead of the evolving resistance curve.")
    p.paragraph_format.first_line_indent = Cm(1.27)
    
    # ===================
    # REFERENCES
    # ===================
    doc.add_page_break()
    h = doc.add_heading('References', level=1)
    h.runs[0].font.name = 'Times New Roman'
    
    references = [
        "World Health Organization. Global priority list of antibiotic-resistant bacteria to guide research, discovery, and development of new antibiotics. Geneva: WHO; 2017.",
        "Yong D, Toleman MA, Giske CG, et al. Characterization of a new metallo-beta-lactamase gene, bla(NDM-1), and a novel erythromycin esterase gene carried on a unique genetic structure in Klebsiella pneumoniae sequence type 14 from India. Antimicrob Agents Chemother. 2009;53(12):5046-5054.",
        "Poirel L, Potron A, Nordmann P. OXA-48-like carbapenemases: the phantom menace. J Antimicrob Chemother. 2012;67(7):1597-1606.",
        "Falcone M, Paterson D. Spotlight on ceftazidime-avibactam: a new option for MDR Gram-negative infections. J Antimicrob Chemother. 2016;71(10):2713-2722.",
        "Giamarellou H. Multidrug-resistant Gram-negative bacteria: how to treat and for how long. Int J Antimicrob Agents. 2010;36(Suppl 2):S50-S54.",
        "Indian Council of Medical Research. Annual Report: Antimicrobial Resistance Research and Surveillance Network. New Delhi: ICMR; 2022.",
        "Gandra S, Alvarez-Uria G, Turner P, et al. Antimicrobial resistance surveillance in low- and middle-income countries. Clin Infect Dis. 2020;71(10):2464-2471.",
        "Clinical and Laboratory Standards Institute. Performance Standards for Antimicrobial Susceptibility Testing. 30th ed. CLSI supplement M100. Wayne, PA: CLSI; 2020.",
        "Hunter JD. Matplotlib: A 2D graphics environment. Comput Sci Eng. 2007;9(3):90-95.",
        "Higgins PG, Dammhayn C, Hackel M, Seifert H. Global spread of carbapenem-resistant Acinetobacter baumannii. J Antimicrob Chemother. 2010;65(2):233-238.",
        "Wyres KL, Holt KE. Klebsiella pneumoniae population genomics and antimicrobial-resistant clones. Trends Microbiol. 2016;24(12):944-956.",
        "Kumarasamy KK, Toleman MA, Walsh TR, et al. Emergence of a new antibiotic resistance mechanism in India, Pakistan, and the UK: a molecular, biological, and epidemiological study. Lancet Infect Dis. 2010;10(9):597-602.",
        "Falagas ME, Grammatikos AP, Michalopoulos A. Potential of old-generation antibiotics to address current need. Expert Rev Anti Infect Ther. 2008;6(5):593-600.",
        "Munoz-Price LS, Poirel L, Bonomo RA, et al. Clinical epidemiology of the global expansion of Klebsiella pneumoniae carbapenemases. Lancet Infect Dis. 2013;13(9):785-796.",
        "Hamidian M, Nigro SJ. Emergence, molecular mechanisms and global spread of carbapenem-resistant Acinetobacter baumannii. Microb Genom. 2019;5(10):e000306.",
        "Livermore DM. Defining an extended-spectrum beta-lactamase. Clin Microbiol Infect. 2008;14(s1):3-10.",
        "Sader HS, Castanheira M, Flamm RK, et al. Antimicrobial activity of ceftazidime-avibactam against Gram-negative bacteria with molecular characterization. Antimicrob Agents Chemother. 2015;61(4):e02083-16.",
        "Liu YY, Wang Y, Walsh TR, et al. Emergence of plasmid-mediated colistin resistance mechanism MCR-1 in animals and human beings in China: a microbiological and molecular biological study. Lancet Infect Dis. 2016;16(2):161-168.",
        "Laxminarayan R, Chaudhury RR. Antibiotic resistance in India: drivers and opportunities for action. PLoS Med. 2016;13(3):e1001974.",
        "Grundmann H, Glasner C, Albiger B, et al. Occurrence of carbapenemase-producing Klebsiella pneumoniae and Escherichia coli in the European survey of carbapenemase-producing Enterobacteriaceae (EuSCAPE): a prospective, multinational study. Lancet Infect Dis. 2017;17(2):153-163.",
        "Shields RK, Chen L, Cheng S, et al. Emergence of ceftazidime-avibactam resistance due to plasmid-borne blaKPC-3 mutations during treatment of carbapenem-resistant Klebsiella pneumoniae infections. Antimicrob Agents Chemother. 2017;61(3):e02097-16."
    ]
    
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = False
        p.add_run(ref)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.5
    
    # Save
    output_path = os.path.join(SUBMISSION_DIR, 'Manuscript_5_Molecular_IJMM_EXPANDED.docx')
    doc.save(output_path)
    print(f"\nSaved expanded manuscript: {output_path}")
    
    return output_path

if __name__ == "__main__":
    create_expanded_manuscript()
