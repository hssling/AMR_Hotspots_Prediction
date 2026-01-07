import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        is_bold = False
        content = part
        if part.startswith('**') and part.endswith('**'):
            is_bold = True
            content = part[2:-2]
        
        sub_parts = re.split(r'(\*.*?\*)', content)
        for sub_part in sub_parts:
            is_italic = False
            sub_content = sub_part
            if sub_part.startswith('*') and sub_part.endswith('*'):
                is_italic = True
                sub_content = sub_part[1:-1]
            
            paragraph.add_run(sub_content).bold = is_bold
            paragraph.runs[-1].italic = is_italic
            paragraph.runs[-1].font.name = 'Times New Roman'
            paragraph.runs[-1].font.size = Pt(12)

def generate_doc(source_md, output_path):
    print(f"Processing {source_md} -> {output_path}")
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    
    with open(source_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            if in_table:
                # Flush Table
                try:
                    valid_rows = [r for r in table_lines if '---' not in r]
                    if valid_rows:
                        cols = len(valid_rows[0].split('|')) - 2
                        tbl = doc.add_table(rows=len(valid_rows), cols=cols)
                        tbl.style = 'Table Grid'
                        for i, row in enumerate(valid_rows):
                            cells = row.split('|')[1:-1]
                            for j, txt in enumerate(cells):
                                if j < len(tbl.columns):
                                    tbl.cell(i, j).text = ""
                                    add_formatted_text(tbl.cell(i, j).paragraphs[0], txt.strip())
                                    if i == 0: tbl.cell(i, j).paragraphs[0].runs[0].font.bold = True
                except: pass
                in_table = False
                table_lines = []
            continue
            
        if line.startswith('## '):
            doc.add_page_break()
            p = doc.add_heading(line.strip('# '), 2)
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.color.rgb = RGBColor(0,0,0)
        elif line.startswith('# '):
            p = doc.add_heading(line.strip('# '), 1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.color.rgb = RGBColor(0,0,0)
        elif line.startswith('![') and '](' in line:
            # Image handling for markdown: ![alt](path)
            try:
                start_path = line.find('](') + 2
                end_path = line.find(')', start_path)
                img_path = line[start_path:end_path]
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(6))
                else:
                    p = doc.add_paragraph(f"[IMAGE MISSING: {img_path}]")
                    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            except Exception as e:
                print(f"Error adding image: {e}")
        elif line.startswith('|'):
            if not in_table: in_table = True; table_lines = [line]
            else: table_lines.append(line)
        else:
            if not in_table:
                p = doc.add_paragraph()
                add_formatted_text(p, line)
            
    doc.save(output_path)

if __name__ == "__main__":
    base_dir = r"C:\Users\hssli\.gemini\antigravity\brain\90c42530-5be5-49cb-a7b5-e960c5582f78"
    out_dir = r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction\submission_rapid"
    os.makedirs(out_dir, exist_ok=True)
    
    # Manuscript 3: TB Triage (Option C)
    generate_doc(os.path.join(base_dir, "Manuscript_3_TB_Triage.md"), os.path.join(out_dir, "Manuscript_3_TB_Triage_Draft.docx"))
    
    # Manuscript 4: AMR Genomics (Option D)
    generate_doc(os.path.join(base_dir, "Manuscript_4_AMR_Genomics.md"), os.path.join(out_dir, "Manuscript_4_AMR_Genomics_Draft.docx"))
