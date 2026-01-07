
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK

def add_formatted_text(paragraph, text):
    """
    Parses text for:
    1. Bold: **text**
    2. Italic: *text*
    3. Citations: $^1$ or [1] -> Superscript
    """
    # 1. Split by Bold
    # regex matches **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        is_bold = False
        content = part
        if part.startswith('**') and part.endswith('**'):
            is_bold = True
            content = part[2:-2]
            
        # 2. Split by Italic (within the chunk)
        # regex matches *text*
        sub_parts = re.split(r'(\*.*?\*)', content)
        
        for sub_part in sub_parts:
            is_italic = False
            sub_content = sub_part
            if sub_part.startswith('*') and sub_part.endswith('*'):
                is_italic = True
                sub_content = sub_part[1:-1]
            
            # 3. Split by Citation (within the sub-chunk)
            # matches $^1$ or [1]
            cit_parts = re.split(r'(\$\^.*?\^?\$|\[\d+(?:[–-]\d+)?\])', sub_content)
            
            for cit_part in cit_parts:
                run = paragraph.add_run()
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = is_bold
                run.italic = is_italic
                
                if (cit_part.startswith('$^') and cit_part.endswith('$')) or \
                   (cit_part.startswith('[') and cit_part.endswith(']') and any(c.isdigit() for c in cit_part)):
                    clean_cit = cit_part.replace('$^', '').replace('$', '').replace('[', '').replace(']', '')
                    run.text = clean_cit
                    run.font.superscript = True
                elif cit_part:
                    # Clean up any leftover markdown symbols if they were somehow missed or nested weirdly
                    # But usually, just raw text
                    run.text = cit_part

def create_manuscript_package(source_md, output_dir, figure_base):
    print(f"Generating Manuscript Package in {output_dir}...")
    os.makedirs(output_dir, exist_ok=True)
    
    # --- Document 1: Main Text (Blinded/Unblinded) ---
    doc_main = Document()
    style = doc_main.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    
    # --- Document 2: Figures Only ---
    doc_figs = Document()
    
    with open(source_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_lines = []
    
    for line in lines:
        line = line.strip()
        if not line: continue
        
        # 1. Page Breaks for Major Sections (Main Text)
        if line.startswith('## Abstract') or line.startswith('## Introduction') or \
           line.startswith('## References') or line.startswith('## Tables') or \
           line.startswith('## Figures'):
            doc_main.add_page_break()
            
        # 2. Headings
        if line.startswith('# '):
            p = doc_main.add_heading(line.strip('# '), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0,0,0)
        elif line.startswith('## '):
            p = doc_main.add_heading(line.strip('# '), level=2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0,0,0)
        elif line.startswith('### '):
            p = doc_main.add_heading(line.strip('# '), level=3)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0,0,0)
            
        # 3. Images -> Go to Figures Doc, Placeholder in Main
        elif line.startswith('![') and '](' in line:
            match = re.search(r'\((.*?)\)', line)
            if match:
                # 1. Clean file URI prefix
                raw_path = match.group(1).replace('file:///', '')
                # 2. Decode URL (fix %20 spaces)
                from urllib.parse import unquote
                rel_path = unquote(raw_path)
                
                if not os.path.isabs(rel_path):
                     # Construct absolute path assuming run from root
                     rel_path = os.path.join(r'd:\research-automation\TB multiomics\AMR_Hotspots_Prediction', rel_path)
                
                # Parsing Caption
                caption = line.split('[')[1].split(']')[0]
                
                # Add to Figures Doc
                if os.path.exists(rel_path):
                    # Add new page for each figure if not first
                    if len(doc_figs.paragraphs) > 0:
                        doc_figs.add_break(WD_BREAK.PAGE)
                        
                    p = doc_figs.add_paragraph()
                    run = p.add_run()
                    try:
                        run.add_picture(rel_path, width=Inches(6.0))
                    except:
                        p.add_run(f"[Could not load image: {rel_path}]")
                    
                    # Caption
                    cap_p = doc_figs.add_paragraph(f"Figure: {caption}")
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add Placeholder to Main Doc
                    p_main = doc_main.add_paragraph()
                    run_main = p_main.add_run(f"[INSERT FIGURE HERE: {caption}]")
                    run_main.bold = True
                else:
                    print(f"Warning: Image not found {rel_path}")

        # 4. Tables
        elif line.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
        
        # 5. Normal Text
        else:
            if in_table:
                # Flush Table to Main Doc
                try:
                    valid_rows = [r for r in table_lines if '---' not in r]
                    if valid_rows:
                        cols = len(valid_rows[0].split('|')) - 2
                        tbl = doc_main.add_table(rows=len(valid_rows), cols=cols)
                        tbl.style = 'Table Grid'
                        for i, row in enumerate(valid_rows):
                            cells = row.split('|')[1:-1]
                            for j, txt in enumerate(cells):
                                # Clean cell text
                                cell_text = txt.strip()
                                # Handle Bold in table
                                tbl.cell(i, j).text = "" # Clear default
                                p_cell = tbl.cell(i, j).paragraphs[0]
                                add_formatted_text(p_cell, cell_text)
                except Exception as e:
                    print(f"Table error: {e}")
                in_table = False
                table_lines = []

            p = doc_main.add_paragraph()
            add_formatted_text(p, line)

    # Save
    main_path = os.path.join(output_dir, "Manuscript_2_IJCCM_Main_Text.docx")
    figs_path = os.path.join(output_dir, "Manuscript_2_IJCCM_Figures.docx")
    
    doc_main.save(main_path)
    doc_figs.save(figs_path)
    print(f"Docs generated: {main_path}, {figs_path}")

if __name__ == "__main__":
    base_md = r"C:\Users\hssli\.gemini\antigravity\brain\90c42530-5be5-49cb-a7b5-e960c5582f78\Manuscript_2_Decoupling_AMR.md"
    out_dir = "submission_manuscript2"
    create_manuscript_package(base_md, out_dir, "outputs/figures_manuscript2")
