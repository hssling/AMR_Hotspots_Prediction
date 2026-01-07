"""
Generate Final Figures Document for Manuscript 5
"""

import os
from docx import Document
from docx.shared import Inches

BASE_DIR = r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction"
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs", "molecular_analysis")
SUBMISSION_DIR = os.path.join(BASE_DIR, "submission_manuscript5")

def create_final_figures_doc():
    print("=" * 60)
    print("GENERATING FINAL FIGURES DOCUMENT")
    print("=" * 60)
    
    doc = Document()
    doc.add_heading('Figures - Manuscript 5', level=1)
    
    # Fig 1
    doc.add_paragraph("Figure 1: Distribution of Resistance Genes by Pathogen")
    fig1 = os.path.join(OUTPUT_DIR, 'fig1_gene_heatmap.png')
    if os.path.exists(fig1):
        print(f"Embedding Figure 1: {fig1}")
        doc.add_picture(fig1, width=Inches(6))
    else:
        print("Warning: Figure 1 not found")
    doc.add_page_break()
    
    # Fig 2
    doc.add_paragraph("Figure 2: Temporal Trends in Resistance Genes")
    fig2 = os.path.join(OUTPUT_DIR, 'fig2_temporal_trends.png')
    if os.path.exists(fig2):
        print(f"Embedding Figure 2: {fig2}")
        doc.add_picture(fig2, width=Inches(6))
    else:
        print("Warning: Figure 2 not found")
    doc.add_page_break()
    
    # Fig 3
    doc.add_paragraph("Figure 3: Susceptibility to Reserve Agents")
    fig3 = os.path.join(OUTPUT_DIR, 'fig3_reserve_agents.png')
    if os.path.exists(fig3):
        print(f"Embedding Figure 3: {fig3}")
        doc.add_picture(fig3, width=Inches(6))
    else:
        print("Warning: Figure 3 not found")
    
    output_path = os.path.join(SUBMISSION_DIR, 'Manuscript_5_Molecular_Figures_FINAL.docx')
    doc.save(output_path)
    print(f"\nSaved Final Figures: {output_path}")

if __name__ == "__main__":
    create_final_figures_doc()
