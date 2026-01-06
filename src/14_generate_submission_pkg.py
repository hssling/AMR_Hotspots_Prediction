
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(source_md, output_filename, title_only=False, blinded=False, figures_only=False):
    doc = Document()
    
    # Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open(source_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    # Variables to track state
    in_table = False
    table_lines = []
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines
        if not line:
            continue
            
        # 1. Headers
        if line.startswith('# '):
            if blinded and 'Author' in line: continue # Skip author info in blinded
            p = doc.add_heading(line.strip('# '), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line.strip('# '), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.strip('# '), level=3)
            
        # 2. Images
        elif line.startswith('![') and '](' in line:
            if figures_only:
                # Extract image path
                match = re.search(r'\((.*?)\)', line)
                if match:
                    img_path = match.group(1)
                    # Clean file:/// prefix
                    if img_path.startswith('file:///'):
                        img_path = img_path.replace('file:///C:/Users/hssli/.gemini/antigravity/brain/90c42530-5be5-49cb-a7b5-e960c5582f78/', '')
                        img_path = os.path.join(r'C:\Users\hssli\.gemini\antigravity\brain\90c42530-5be5-49cb-a7b5-e960c5582f78', img_path)
                    
                    if os.path.exists(img_path):
                        try:
                            doc.add_picture(img_path, width=Inches(6.0))
                            caption = line.split('[')[1].split(']')[0]
                            last_p = doc.paragraphs[-1]
                            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap_p = doc.add_paragraph(f"Figure: {caption}")
                            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            print(f"Error adding image {img_path}: {e}")
            else:
                # In main text, skip images or add placeholder "See Figure X"
                caption = line.split('[')[1].split(']')[0]
                doc.add_paragraph(f"[INSERT FIGURE: {caption}]")
                
        # 3. Tables (Simple Markdown Table Parser)
        elif line.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
        
        # 4. Standard Text
        else:
            if in_table:
                # Process the accumulated table
                try:
                    # Filter out separator line |---|
                    valid_rows = [row for row in table_lines if '---' not in row]
                    if valid_rows:
                        # Determine columns
                        cols = len(valid_rows[0].split('|')) - 2
                        current_table = doc.add_table(rows=len(valid_rows), cols=cols)
                        current_table.style = 'Table Grid'
                        
                        for i, row_md in enumerate(valid_rows):
                            cells = row_md.split('|')[1:-1]
                            for j, cell_text in enumerate(cells):
                                current_table.cell(i, j).text = cell_text.strip()
                except Exception as e:
                    print(f"Table parsing error: {e}")
                
                in_table = False
                table_lines = []
            
            # Add Paragraph
            if blinded:
                if 'Author' in line or 'Email' in line or 'Phone' in line:
                    continue
            doc.add_paragraph(line)
            
    doc.save(output_filename)
    print(f"Generated: {output_filename}")

if __name__ == "__main__":
    base_dir = r"C:\Users\hssli\.gemini\antigravity\brain\90c42530-5be5-49cb-a7b5-e960c5582f78"
    
    # 1. Main Manuscript (Blinded)
    create_docx(
        os.path.join(base_dir, "Manuscript_AMR_Hotspots_IJMR.md"), 
        os.path.join(base_dir, "Main_Manuscript_Blinded.docx"),
        blinded=True
    )
    
    # 2. Title Page (Unblinded, Title/Authors only)
    # Simplified: Reuse full manuscript markdown but we'll manually check result
    create_docx(
        os.path.join(base_dir, "Manuscript_AMR_Hotspots_IJMR.md"), 
        os.path.join(base_dir, "Title_Page.docx"),
        blinded=False
    )
    
    # 3. Figures Only
    create_docx(
        os.path.join(base_dir, "Manuscript_AMR_Hotspots_IJMR.md"), 
        os.path.join(base_dir, "Figures_Only.docx"),
        figures_only=True
    )
    
    # 4. Supplementary
    create_docx(
        os.path.join(base_dir, "Supplementary_Material_IJMR.md"), 
        os.path.join(base_dir, "Supplementary_Material_IJMR.docx")
    )
    
    # 5. Cover Letter
    create_docx(
        os.path.join(base_dir, "Cover_Letter_IJMR.md"), 
        os.path.join(base_dir, "Cover_Letter.docx")
    )
