"""
Generate Cover Letter DOCX for IJMR Submission
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cover_letter():
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Date
    p = doc.add_paragraph("January 07, 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # To
    doc.add_paragraph("To,")
    doc.add_paragraph("The Editor-in-Chief")
    doc.add_paragraph("Indian Journal of Medical Research (IJMR)")
    doc.add_paragraph("Publication Division, Indian Council of Medical Research")
    doc.add_paragraph("New Delhi, India")
    
    doc.add_paragraph()
    
    # Subject
    p = doc.add_paragraph()
    p.add_run("Subject: ").bold = True
    p.add_run("Submission of Original Article – \"Spatiotemporal Modeling of Antimicrobial Resistance Hotspots in India (2017-2024)\"")
    
    doc.add_paragraph()
    
    # Body
    doc.add_paragraph("Dear Editor,")
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("We are pleased to submit our original research article titled ")
    p.add_run("\"Spatiotemporal Modeling of Antimicrobial Resistance Hotspots in India (2017-2024): Integrating Genomic Surveillance with Longitudinal Predictive Analytics\"").bold = True
    p.add_run(" for consideration for publication in the Indian Journal of Medical Research.")
    
    doc.add_paragraph()
    
    # Relevance Section
    p = doc.add_paragraph()
    p.add_run("Relevance to IJMR").bold = True
    
    doc.add_paragraph("This study directly leverages data from the ICMR-AMRSN network—the very surveillance system established and maintained by ICMR—making IJMR the ideal venue for dissemination. Our findings provide actionable intelligence for:")
    
    doc.add_paragraph("ICMR Policy Planning: Identification of NDM-dominated Northern and OXA-dominated Southern resistance corridors.", style='List Bullet')
    doc.add_paragraph("NAP-AMR Implementation: A predictive framework to transition surveillance from retrospective to proactive.", style='List Bullet')
    doc.add_paragraph("National Stewardship: Evidence supporting hospital-level antibiograms and hyper-local interventions.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Key Findings
    p = doc.add_paragraph()
    p.add_run("Key Findings").bold = True
    
    doc.add_paragraph("Carbapenem resistance in K. pneumoniae rose from 41.5% (2017) to >57% (2021).", style='List Bullet')
    doc.add_paragraph("Machine learning model achieved R² = 0.87 for predicting resistance trajectories.", style='List Bullet')
    doc.add_paragraph("Distinct geo-genomic hotspots identified: NDM (North) vs. OXA-23 (South).", style='List Bullet')
    doc.add_paragraph("Positive correlation (r=0.10) between resistance rates and ICU mortality.", style='List Bullet')
    
    doc.add_paragraph()
    
    # Manuscript Details
    p = doc.add_paragraph()
    p.add_run("Manuscript Details").bold = True
    
    doc.add_paragraph("Word Count: ~3,500 words (Main Text)", style='List Bullet')
    doc.add_paragraph("Tables: 2", style='List Bullet')
    doc.add_paragraph("Figures: 5", style='List Bullet')
    doc.add_paragraph("References: 30", style='List Bullet')
    
    doc.add_paragraph()
    
    # Declarations
    p = doc.add_paragraph()
    p.add_run("Declarations").bold = True
    
    doc.add_paragraph("This manuscript is original and has not been published previously nor is under consideration elsewhere.", style='List Bullet')
    doc.add_paragraph("All authors have read and approved the final manuscript.", style='List Bullet')
    doc.add_paragraph("Ethical approval was waived as only anonymized, aggregated secondary data from public domain ICMR reports was used.", style='List Bullet')
    doc.add_paragraph("The authors declare no conflicts of interest.", style='List Bullet')
    doc.add_paragraph("This research was self-funded.", style='List Bullet')
    doc.add_paragraph("AI tools (Gemini Labs) were used for data extraction and drafting assistance, disclosed per ICMJE guidelines.", style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_paragraph("We believe this work aligns with IJMR's mission to advance medical research in India and provides a template for predictive AMR surveillance.")
    
    doc.add_paragraph()
    doc.add_paragraph("Thank you for considering our submission. We look forward to your favorable response.")
    
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    
    # Signature
    p = doc.add_paragraph()
    p.add_run("Dr. Siddalingaiah H S").bold = True
    doc.add_paragraph("Professor, Department of Community Medicine")
    doc.add_paragraph("Shridevi Institute of Medical Sciences and Research Hospital")
    doc.add_paragraph("Tumkur, Karnataka, India – 572106")
    doc.add_paragraph("Email: hssling@yahoo.com")
    doc.add_paragraph("Phone: +91-8941087719")
    doc.add_paragraph("ORCID: 0000-0002-4771-8285")
    
    # Save
    out_path = r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction\submission\Cover_Letter_IJMR_V2.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    create_cover_letter()

