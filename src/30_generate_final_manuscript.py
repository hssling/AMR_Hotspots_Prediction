"""
Robust Manuscript Generator for IJMR Submission (V3 - Final)
- Clean citation format: [N] or [N,M] -> Word superscript
- No LaTeX/$ symbols
- Proper Bold/Italic handling
- Embedded images
- Tables at end
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_formatted_text(paragraph, text):
    """
    Parses markdown-like text and adds it to a Word paragraph with proper formatting.
    Handles: **bold**, *italic*, and [N] or [N,M] citations as superscripts.
    """
    if not text:
        return

    # Pattern to split by: **bold**, *italic*, and [citations]
    # We process in order: Bold first, then Italic, then Citations
    
    # Step 1: Split by Bold (**...**)
    bold_pattern = r'(\*\*.*?\*\*)'
    tokens = re.split(bold_pattern, text)
    
    for token in tokens:
        if not token:
            continue
            
        is_bold = False
        content = token
        
        if token.startswith('**') and token.endswith('**'):
            is_bold = True
            content = token[2:-2]
        
        # Step 2: Split by Italic (*...* but not **)
        # Need to be careful not to match inside ** 
        italic_pattern = r'(?<!\*)(\*[^*]+?\*)(?!\*)'
        sub_tokens = re.split(italic_pattern, content)
        
        for sub_token in sub_tokens:
            if not sub_token:
                continue
                
            is_italic = False
            sub_content = sub_token
            
            if sub_token.startswith('*') and sub_token.endswith('*') and not sub_token.startswith('**'):
                is_italic = True
                sub_content = sub_token[1:-1]
            
            # Step 3: Split by Citations [N] or [N,M]
            # Pattern: [1] or [1,2] or [1-3] or [1,2,10]
            citation_pattern = r'(\[\d+(?:[,\-–]\d+)*\])'
            atoms = re.split(citation_pattern, sub_content)
            
            for atom in atoms:
                if not atom:
                    continue
                    
                is_citation = False
                atom_text = atom
                
                if re.match(r'^\[\d+(?:[,\-–]\d+)*\]$', atom):
                    is_citation = True
                    atom_text = atom[1:-1]  # Remove [ and ]
                
                run = paragraph.add_run(atom_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
                if is_bold:
                    run.bold = True
                if is_italic:
                    run.italic = True
                if is_citation:
                    run.font.superscript = True

def create_manuscript(source_md, output_dir, figure_base):
    """Main function to convert markdown to DOCX."""
    print("=" * 60)
    print("GENERATING PUBLICATION-READY MANUSCRIPT (V3)")
    print("=" * 60)
    
    doc = Document()
    
    # Set default styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    
    with open(source_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    tables_buffer = []
    current_table = []
    in_table = False
    table_title_buffer = None
    
    for line in lines:
        line = line.rstrip('\r\n')
        stripped = line.strip()
        
        # --- Table Handling (Capture for end of document) ---
        if stripped.startswith('**Table'):
            table_title_buffer = stripped
            continue
        
        if stripped.startswith('|'):
            in_table = True
            current_table.append(stripped)
            continue
        
        if in_table and not stripped.startswith('|'):
            if current_table:
                tables_buffer.append({'title': table_title_buffer, 'rows': current_table})
            current_table = []
            in_table = False
            table_title_buffer = None
            if not stripped:
                continue
        
        # --- Skip empty lines (add paragraph break) ---
        if not stripped:
            if len(doc.paragraphs) > 0:
                doc.add_paragraph()
            continue
        
        # --- Skip horizontal rules ---
        if stripped == '---':
            continue
        
        # --- Headers ---
        if stripped.startswith('## '):
            doc.add_page_break()
            heading_text = stripped[3:]
            p = doc.add_heading(heading_text, level=2)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.color.rgb = RGBColor(0, 0, 0)
            continue
        
        if stripped.startswith('# '):
            heading_text = stripped[2:]
            p = doc.add_heading(heading_text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.color.rgb = RGBColor(0, 0, 0)
            continue
        
        if stripped.startswith('### '):
            heading_text = stripped[4:]
            p = doc.add_heading(heading_text, level=3)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.color.rgb = RGBColor(0, 0, 0)
            continue
        
        # --- Images ---
        if stripped.startswith('![') and '](' in stripped:
            try:
                start = stripped.find('](') + 2
                end = stripped.find(')', start)
                img_path = stripped[start:end]
                
                # Resolve path
                search_dirs = [
                    figure_base,
                    r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction\outputs\figures"
                ]
                
                final_path = None
                for search_dir in search_dirs:
                    candidate = os.path.join(search_dir, os.path.basename(img_path))
                    if os.path.exists(candidate):
                        final_path = candidate
                        break
                
                if final_path:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(final_path, width=Inches(5.5))
                    print(f"  [OK] Embedded: {os.path.basename(final_path)}")
                else:
                    print(f"  [WARN] Image not found: {img_path}")
                    p = doc.add_paragraph(f"[IMAGE NOT FOUND: {img_path}]")
                    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            except Exception as e:
                print(f"  [ERROR] Image processing: {e}")
            continue
        
        # --- Figure Captions (lines starting with *Figure) ---
        if stripped.startswith('*Figure') or stripped.startswith('*Table'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Remove leading/trailing * for italic
            caption_text = stripped.strip('*')
            run = p.add_run(caption_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.italic = True
            continue
        
        # --- Regular Text Paragraphs ---
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
    
    # --- Flush any remaining table ---
    if in_table and current_table:
        tables_buffer.append({'title': table_title_buffer, 'rows': current_table})
    
    # --- Append Tables at the END ---
    if tables_buffer:
        doc.add_page_break()
        h = doc.add_heading("Tables", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in h.runs:
            r.font.name = 'Times New Roman'
            r.font.color.rgb = RGBColor(0, 0, 0)
        
        for tbl in tables_buffer:
            doc.add_paragraph()
            
            # Table Title
            if tbl['title']:
                p = doc.add_paragraph()
                add_formatted_text(p, tbl['title'])
            
            # Table Data (skip header separator row)
            data_rows = [r for r in tbl['rows'] if ':---' not in r and '---:' not in r and '| :---' not in r]
            
            if data_rows:
                num_cols = len(data_rows[0].split('|')) - 2
                table = doc.add_table(rows=len(data_rows), cols=max(1, num_cols))
                table.style = 'Table Grid'
                
                for i, row_str in enumerate(data_rows):
                    cells = row_str.split('|')[1:-1]
                    for j, cell_text in enumerate(cells):
                        if j < len(table.columns):
                            cell_para = table.cell(i, j).paragraphs[0]
                            add_formatted_text(cell_para, cell_text.strip())
                            # Bold header row
                            if i == 0:
                                for run in cell_para.runs:
                                    run.bold = True
    
    # --- Save ---
    out_file = os.path.join(output_dir, "Manuscript_1_IJMR_FINAL_V3.docx")
    doc.save(out_file)
    print("=" * 60)
    print(f"SUCCESS: Saved to {out_file}")
    print("=" * 60)

if __name__ == "__main__":
    base_dir = r"C:\Users\hssli\.gemini\antigravity\brain\90c42530-5be5-49cb-a7b5-e960c5582f78"
    out_dir = r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction\submission"
    create_manuscript(
        os.path.join(base_dir, "Manuscript_AMR_Hotspots_IJMR_CLEAN.md"),
        out_dir,
        base_dir
    )
