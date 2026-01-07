"""
Generate Final Tables Document for Manuscript 5
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Pt

BASE_DIR = r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction"
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs", "molecular_analysis")
SUBMISSION_DIR = os.path.join(BASE_DIR, "submission_manuscript5")

def create_final_tables_doc():
    print("=" * 60)
    print("GENERATING FINAL TABLES DOCUMENT")
    print("=" * 60)
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    doc.add_heading('Tables - Manuscript 5', level=1)
    
    # Table 1
    doc.add_paragraph("Table 1: Prevalence of Resistance Genes by Pathogen")
    csv_path1 = os.path.join(OUTPUT_DIR, 'table1_gene_prevalence.csv')
    if os.path.exists(csv_path1):
        t1 = pd.read_csv(csv_path1)
        table = doc.add_table(rows=len(t1)+1, cols=len(t1.columns))
        table.style = 'Table Grid'
        
        # Header
        for j, col in enumerate(t1.columns):
            cell = table.cell(0, j)
            cell.text = col
            cell.paragraphs[0].runs[0].font.bold = True
            
        # Rows
        for i, row in t1.iterrows():
            for j, val in enumerate(row):
                table.cell(i+1, j).text = str(val) if pd.notna(val) else "-"
        print(f"Added Table 1 from {csv_path1}")
    else:
        print("Warning: Table 1 CSV not found")
        
    doc.add_page_break()
    
    # Table 2
    doc.add_paragraph("Table 2: Susceptibility to Reserve Antimicrobial Agents")
    csv_path2 = os.path.join(OUTPUT_DIR, 'table2_reserve_susceptibility.csv')
    if os.path.exists(csv_path2):
        t2 = pd.read_csv(csv_path2)
        table = doc.add_table(rows=len(t2)+1, cols=len(t2.columns))
        table.style = 'Table Grid'
        
        # Header
        for j, col in enumerate(t2.columns):
            cell = table.cell(0, j)
            cell.text = col
            cell.paragraphs[0].runs[0].font.bold = True
            
        # Rows
        for i, row in t2.iterrows():
            for j, val in enumerate(row):
                table.cell(i+1, j).text = str(val) if pd.notna(val) else "-"
        print(f"Added Table 2 from {csv_path2}")
    else:
        print("Warning: Table 2 CSV not found")
    
    output_path = os.path.join(SUBMISSION_DIR, 'Manuscript_5_Molecular_Tables_FINAL.docx')
    doc.save(output_path)
    print(f"\nSaved Final Tables: {output_path}")

if __name__ == "__main__":
    create_final_tables_doc()
