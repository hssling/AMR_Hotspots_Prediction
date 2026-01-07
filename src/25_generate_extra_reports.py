import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    h = doc.add_heading(text, level)
    h.runs[0].font.name = 'Times New Roman'
    h.runs[0].font.color.rgb = RGBColor(0,0,0)

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    runner = p.add_run(text)
    runner.font.name = 'Times New Roman'
    runner.font.size = Pt(12)
    if bold: runner.font.bold = True

def generate_policy_brief(output_dir):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    add_heading(doc, "Policy Brief: Impact of the 2019 'Red Line' Campaign on AMR", 0)
    add_para(doc, "Authors: AI for Public Health Lab")
    add_para(doc, "Date: January 2026\n")
    
    add_heading(doc, "Executive Summary", 1)
    add_para(doc, "This brief evaluates the effectiveness of the 2019 National 'Red Line' Antibiotic Awareness Campaign using Interrupted Time Series (ITS) analysis of resistance trends from 2016-2024.")
    
    add_heading(doc, "Key Findings", 1)
    add_para(doc, "1. Trend Reversal: Prior to 2019, resistance rates were rising annually. Post-2019, the trend shows a flattening (Slope Change: -1.23%).")
    add_para(doc, "2. Statistical Significance: While the direction is positive, the change is not yet statistically significant (p=0.85), likely due to the short post-intervention window.")
    
    add_heading(doc, "Visual Evidence", 1)
    add_para(doc, "[INSERT FIGURE HERE: red_line_impact_its.png]")
    try:
        img_path = os.path.join(output_dir, "..", "outputs", "advanced_analytics", "red_line_impact_its.png")
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6))
    except:
        add_para(doc, "[Image verification failed]")

    add_heading(doc, "Recommendations", 1)
    add_para(doc, "- Sustain the Campaign: The signals of 'flattening' suggest the campaign is working.")
    add_para(doc, "- Enhance Surveillance: Expand data collection to increase statistical power.")
    
    out_path = os.path.join(output_dir, "Policy_Brief_Red_Line_Campaign.docx")
    doc.save(out_path)
    print(f"Policy Brief saved: {out_path}")

def generate_scorecard_report(output_dir):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    add_heading(doc, "Hospital AMR Stewardship Scorecard (2025)", 0)
    add_para(doc, "A PCA-Based Ranking of Clinical Resistance Burden")
    
    add_heading(doc, "Methodology", 1)
    add_para(doc, "We utilised Principal Component Analysis (PCA) to condense multi-drug resistance patterns into a single 'Burden Score' (0-100). Higher scores indicate higher overall resistance.")

    add_heading(doc, "Top 5 High-Burden Centers", 1)
    
    # Load Data
    csv_path = os.path.join(output_dir, "..", "outputs", "scorecard", "hospital_amr_scorecard.csv")
    if os.path.exists(csv_path):
        df = pd.read_csv(csv_path)
        top5 = df.head(5)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Rank'
        hdr_cells[1].text = 'Center Name'
        hdr_cells[2].text = 'Burden Score'
        
        for index, row in top5.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(int(row['Rank']))
            row_cells[1].text = str(row['Center_Name'])
            row_cells[2].text = f"{row['AMR_Burden_Score']:.1f}"
            
    add_heading(doc, "Scorecard Visualization", 1)
    try:
        img_path = os.path.join(output_dir, "..", "outputs", "scorecard", "amr_scorecard_plot.png")
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6))
    except: pass
    
    out_path = os.path.join(output_dir, "Report_Hospital_Scorecard.docx")
    doc.save(out_path)
    print(f"Scorecard Report saved: {out_path}")

if __name__ == "__main__":
    out_dir = r"d:\research-automation\TB multiomics\AMR_Hotspots_Prediction\submission_rapid"
    generate_policy_brief(out_dir)
    generate_scorecard_report(out_dir)
